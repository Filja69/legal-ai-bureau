"""DOCX extraction — paragraphs, headings, tables (Phase 9.2 brief §3/§11).

`app/documents/validation.py` already ran the ZIP-bomb/manifest checks
before this is ever called (brief §6) — this module trusts its caller.
"""
from __future__ import annotations

import zipfile
from io import BytesIO

import docx
from docx.opc.exceptions import PackageNotFoundError

from app.documents.extraction.base import ExtractedDocument, ExtractedSection, ExtractedTable, ExtractionError


class DocxExtractor:
    async def extract(self, content: bytes) -> ExtractedDocument:
        try:
            document = docx.Document(BytesIO(content))
        except (PackageNotFoundError, KeyError, ValueError, zipfile.BadZipFile) as exc:
            raise ExtractionError("CORRUPTED_FILE", "DOCX could not be parsed — the file may be corrupted") from exc

        sections: list[ExtractedSection] = []
        heading_stack: list[str] = []  # current heading path, e.g. ["Раздел 4", "4.2"]

        for paragraph in document.paragraphs:
            text = paragraph.text
            if not text.strip():
                continue
            style_name = paragraph.style.name if paragraph.style is not None else None
            is_heading = style_name is not None and style_name.lower().startswith("heading")
            if is_heading:
                assert style_name is not None  # narrowed by is_heading above
                # A heading at level N replaces anything below level N in the
                # stack, so section_path always reflects genuine document
                # nesting rather than accumulating unrelated headings.
                try:
                    level = int(style_name.rsplit(" ", 1)[-1])
                except ValueError:
                    level = 1
                heading_stack = heading_stack[: level - 1] + [text.strip()]
            section_path = " > ".join(heading_stack) if heading_stack else None
            sections.append(ExtractedSection(text=text, section_path=section_path, is_heading=is_heading))

        tables: list[ExtractedTable] = []
        for table in document.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(ExtractedTable(page_number=None, rows=rows))

        full_text = "\n".join(s.text for s in sections)
        if tables:
            table_text = "\n\n".join(
                "\n".join(" | ".join(cell for cell in row) for row in table.rows) for table in tables
            )
            full_text = f"{full_text}\n\n{table_text}"

        return ExtractedDocument(
            text=full_text,
            sections=sections,
            tables=tables,
            metadata={"paragraph_count": len(sections), "table_count": len(tables)},
            extractor="python-docx",
        )
