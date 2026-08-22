"""Documents — LEGAL-API.md §Documents. Phase 8 added GET /documents (list);
Phase 9.2 added real upload validation + synchronous processing (upload now
returns the document already processed, not merely "uploaded" — see
tests/integration/test_document_pipeline.py for the full pipeline suite).
"""
from __future__ import annotations

import io

import docx
import pytest
import sqlalchemy as sa

from app.documents.storage.base import DocumentStorageConfigError, DocumentStorageError
from app.models.matters import DocumentChunk, DocumentStatus
from app.rag.embeddings.base import EmbeddingProviderError
from tests.security.auth_factories import make_org_and_workspace


def _synthetic_loan_docx() -> bytes:
    """A real .docx (not a fixture file) — matches the P0 regression
    scenario exactly: a normal Word document with several paragraphs, none
    of which is malformed. Building it with python-docx proves the upload
    path against a genuinely valid Office document, not a hand-crafted
    ZIP that merely satisfies validate_upload()'s checks.
    """
    document = docx.Document()
    document.add_paragraph("ДОГОВОР ЗАЙМА")
    document.add_paragraph("Сумма займа составляет 5 000 000 рублей.")
    document.add_paragraph("Заемщик обязан возвратить сумму займа.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.asyncio
async def test_list_documents_empty_for_new_workspace(client, db_session):
    _org, workspace = await make_org_and_workspace(db_session)
    await db_session.commit()
    response = await client.get("/api/v1/legal/documents", headers={"X-Workspace-Id": str(workspace.id)})
    assert response.status_code == 200
    assert response.json() == []


@pytest.mark.asyncio
async def test_upload_then_list_shows_the_document(client, db_session):
    _org, workspace = await make_org_and_workspace(db_session)
    await db_session.commit()
    headers = {"X-Workspace-Id": str(workspace.id)}

    files = {"file": ("contract.txt", io.BytesIO(b"hello world"), "text/plain")}
    upload_response = await client.post("/api/v1/legal/documents", files=files, headers=headers)
    assert upload_response.status_code == 201
    uploaded = upload_response.json()
    assert uploaded["status"] == "ready"  # Phase 9.2 — upload now synchronously processes to completion
    assert uploaded["title"] == "contract.txt"

    list_response = await client.get("/api/v1/legal/documents", headers=headers)
    assert list_response.status_code == 200
    items = list_response.json()
    assert len(items) == 1
    assert items[0]["id"] == uploaded["id"]


@pytest.mark.asyncio
async def test_list_documents_is_workspace_isolated(client, db_session):
    _org_a, workspace_a = await make_org_and_workspace(db_session, "Docs Org A")
    _org_b, workspace_b = await make_org_and_workspace(db_session, "Docs Org B")
    await db_session.commit()

    files = {"file": ("secret.txt", io.BytesIO(b"secret"), "text/plain")}
    await client.post("/api/v1/legal/documents", files=files, headers={"X-Workspace-Id": str(workspace_a.id)})

    other_workspace_list = await client.get("/api/v1/legal/documents", headers={"X-Workspace-Id": str(workspace_b.id)})
    assert other_workspace_list.json() == []


# --- P0 production incident regression (DOCX upload appeared to fail with a
# misleading browser "CORS blocked" error; real root cause: an unhandled
# exception in storage.put()/embedding.embed() escaped past CORSMiddleware —
# see app/documents/storage/base.py's DocumentStorageError docstring for the
# full Starlette-middleware-ordering explanation). ---


@pytest.mark.asyncio
async def test_real_docx_upload_completes_full_pipeline_to_ready(client, db_session):
    """The exact end-to-end path the real lawyer's browser exercises:
    upload -> accepted -> Document created -> text extracted -> content
    non-empty -> status READY -> chunks/indexing created. A genuinely valid
    .docx (per _synthetic_loan_docx) was never the problem — this proves it.
    """
    _org, workspace = await make_org_and_workspace(db_session)
    await db_session.commit()
    headers = {"X-Workspace-Id": str(workspace.id)}

    files = {
        "file": (
            "loan_agreement.docx",
            io.BytesIO(_synthetic_loan_docx()),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    response = await client.post("/api/v1/legal/documents", files=files, headers=headers)

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "ready"
    assert body["title"] == "loan_agreement.docx"

    document_id = body["id"]
    text_response = await client.get(f"/api/v1/legal/documents/{document_id}/text", headers=headers)
    assert text_response.status_code == 200
    extracted_text = text_response.json()["text"]
    assert extracted_text.strip() != ""
    assert "ДОГОВОР ЗАЙМА" in extracted_text
    assert "5 000 000" in extracted_text

    chunks = (await db_session.execute(sa.select(DocumentChunk).where(DocumentChunk.document_id == document_id))).scalars().all()
    assert len(chunks) > 0


@pytest.mark.asyncio
async def test_real_docx_upload_works_with_configured_persistent_volume_path(client, db_session, monkeypatch, tmp_path):
    """Same full pipeline as test_real_docx_upload_completes_full_pipeline_to_ready,
    but with LOCAL_STORAGE_PATH-style configuration active (the actual
    Railway persistent-Volume scenario) instead of the repo-relative dev
    default — proves the DOCX pipeline doesn't care which root is configured.
    """
    import app.documents.storage.local_storage as local_storage_module

    monkeypatch.setattr(local_storage_module, "_STORAGE_ROOT", tmp_path / "persistent-volume-mount")

    _org, workspace = await make_org_and_workspace(db_session)
    await db_session.commit()
    headers = {"X-Workspace-Id": str(workspace.id)}

    files = {
        "file": (
            "loan_agreement.docx",
            io.BytesIO(_synthetic_loan_docx()),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    response = await client.post("/api/v1/legal/documents", files=files, headers=headers)

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "ready"

    text_response = await client.get(f"/api/v1/legal/documents/{body['id']}/text", headers=headers)
    assert "ДОГОВОР ЗАЙМА" in text_response.json()["text"]


@pytest.mark.asyncio
async def test_storage_failure_returns_clean_503_with_cors_headers_not_generic_500(client, db_session, monkeypatch):
    """Reproduces the exact P0 mechanism: storage.put() raises. Before the
    fix this was an unhandled exception that produced a response with no
    Access-Control-Allow-Origin header — indistinguishable, from the
    browser's perspective, from a genuine CORS misconfiguration. After the
    fix it's a clean 503 that DOES carry CORS headers, because it flows
    through ExceptionMiddleware -> CORSMiddleware like any other HTTPException.
    """
    _org, workspace = await make_org_and_workspace(db_session)
    await db_session.commit()

    class _AlwaysFailsStorage:
        async def put(self, *args, **kwargs):
            raise DocumentStorageError("S3 put_object failed (ClientError)")

    import app.api.v1.documents as documents_module

    monkeypatch.setattr(documents_module, "get_document_storage", lambda: _AlwaysFailsStorage())

    files = {"file": ("contract.txt", io.BytesIO(b"hello world"), "text/plain")}
    response = await client.post(
        "/api/v1/legal/documents",
        files=files,
        headers={"X-Workspace-Id": str(workspace.id), "Origin": "http://localhost:3000"},
    )

    assert response.status_code == 503
    detail = response.json()["detail"]
    assert detail  # a real, specific, safe message — never the generic frontend fallback
    assert "недоступен" in detail.lower()
    assert "traceback" not in detail.lower() and "clienterror" not in detail.lower()  # no raw exception leaked
    assert response.headers.get("access-control-allow-origin") == "http://localhost:3000"  # THE regression check

    # No Document row was created — storage.put() fails before the row exists.
    list_response = await client.get("/api/v1/legal/documents", headers={"X-Workspace-Id": str(workspace.id)})
    assert list_response.json() == []


@pytest.mark.asyncio
async def test_storage_config_error_returns_clean_503_with_cors_headers(client, db_session, monkeypatch):
    """The EXACT real production traceback: get_document_storage() itself
    raises DocumentStorageConfigError (STORAGE_PROVIDER=s3 with no
    STORAGE_BUCKET) while constructing the backend, before .put() is even
    reached. Previously only DocumentStorageError (a sibling, not a parent,
    class) was caught here — this exact exception type was still escaping
    uncaught. Must produce the same clean, CORS-intact 503 as a runtime
    storage failure, not the original CORS-masking bug.
    """
    _org, workspace = await make_org_and_workspace(db_session)
    await db_session.commit()

    def _broken_factory():
        raise DocumentStorageConfigError("STORAGE_PROVIDER=s3 but STORAGE_BUCKET is not set.")

    import app.api.v1.documents as documents_module

    monkeypatch.setattr(documents_module, "get_document_storage", _broken_factory)

    files = {"file": ("contract.txt", io.BytesIO(b"hello world"), "text/plain")}
    response = await client.post(
        "/api/v1/legal/documents",
        files=files,
        headers={"X-Workspace-Id": str(workspace.id), "Origin": "http://localhost:3000"},
    )

    assert response.status_code == 503
    assert "STORAGE_BUCKET" not in response.json()["detail"]  # no raw config detail leaked to the client
    assert response.headers.get("access-control-allow-origin") == "http://localhost:3000"


@pytest.mark.asyncio
async def test_reprocess_storage_failure_returns_clean_503_not_generic_500(client, db_session, monkeypatch):
    """Sibling gap to test_storage_failure_returns_clean_503_with_cors_headers_not_generic_500
    above, found while validating OCR reprocessing of already-uploaded
    documents: POST /documents/{id}/process's own storage.get() call was
    never guarded the same way upload's storage.put() was — an unguarded
    exception here hits the exact same CORS-masking mechanism.
    """
    _org, workspace = await make_org_and_workspace(db_session)
    await db_session.commit()
    headers = {"X-Workspace-Id": str(workspace.id)}

    upload_response = await client.post(
        "/api/v1/legal/documents", files={"file": ("contract.txt", io.BytesIO(b"hello world"), "text/plain")}, headers=headers
    )
    document_id = upload_response.json()["id"]

    class _AlwaysFailsStorage:
        async def get(self, *args, **kwargs):
            raise DocumentStorageError("local disk read failed (OSError)")

    import app.api.v1.documents as documents_module

    monkeypatch.setattr(documents_module, "get_document_storage", lambda: _AlwaysFailsStorage())

    response = await client.post(
        f"/api/v1/legal/documents/{document_id}/process", headers={**headers, "Origin": "http://localhost:3000"}
    )

    assert response.status_code == 503
    detail = response.json()["detail"]
    assert "недоступен" in detail.lower()
    assert "traceback" not in detail.lower() and "oserror" not in detail.lower()
    assert response.headers.get("access-control-allow-origin") == "http://localhost:3000"


@pytest.mark.asyncio
async def test_embedding_failure_marks_document_failed_not_stuck_or_500(client, db_session, monkeypatch):
    """Same P0 mechanism, later in the pipeline: embedding.embed() raises
    after the Document row already exists. Must not: crash the request (no
    500), leave status stuck at PROCESSING, or create partial DocumentChunk
    rows. Must: return 201 (the upload itself DID succeed), with the
    document honestly marked FAILED and a real reason — same discipline
    already used for extraction failures.
    """
    _org, workspace = await make_org_and_workspace(db_session)
    await db_session.commit()

    class _AlwaysFailsEmbeddingProvider:
        model_name = "test-model"
        dimensions = 8
        provider_name = "test"
        model_version = None

        async def embed(self, texts):
            raise EmbeddingProviderError("connection to embedding provider timed out")

    import app.domains.documents.pipeline as pipeline_module

    monkeypatch.setattr(pipeline_module, "get_embedding_provider", lambda: _AlwaysFailsEmbeddingProvider())

    files = {"file": ("contract.txt", io.BytesIO(b"Some real extractable text content."), "text/plain")}
    response = await client.post(
        "/api/v1/legal/documents", files=files, headers={"X-Workspace-Id": str(workspace.id)}
    )

    assert response.status_code == 201  # the upload itself succeeded — this is a processing outcome, not a request failure
    body = response.json()
    assert body["status"] == "failed"
    assert body["processing_error"]
    assert "traceback" not in body["processing_error"].lower()

    chunks = (await db_session.execute(sa.select(DocumentChunk).where(DocumentChunk.document_id == body["id"]))).scalars().all()
    assert chunks == []  # no partial/half-ready chunks

    from app.models.matters import Document

    document = await db_session.get(Document, body["id"])
    assert document.status == DocumentStatus.FAILED  # never left at PROCESSING
