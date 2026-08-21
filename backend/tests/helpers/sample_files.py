"""Byte-level fixtures for document-intelligence tests — real PDF/DOCX/XLSX
bytes built programmatically (not fixture files on disk), so the test suite
has no external file dependencies.
"""
from __future__ import annotations

import zipfile
from io import BytesIO


def build_minimal_pdf(text: str, *, page_count: int = 1) -> bytes:
    """A hand-assembled, spec-valid single/multi-page PDF with one text
    string per page — good enough for `pypdf` to parse and extract real
    text from, without depending on a PDF-generation library.

    `text` must be Latin-1-representable: the minimal font resource here is
    the standard Helvetica font with no custom encoding, which (like real
    PDFs using WinAnsiEncoding) can't represent Cyrillic — a real product
    PDF with Cyrillic text embeds a font with a different encoding, which
    is out of scope for a byte-level test fixture.
    """
    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")

    page_obj_ids = list(range(3, 3 + page_count))
    kids = " ".join(f"{pid} 0 R" for pid in page_obj_ids)
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>".encode())

    font_obj_id = 3 + page_count
    content_obj_ids = list(range(font_obj_id + 1, font_obj_id + 1 + page_count))
    for i in range(page_count):
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 {font_obj_id} 0 R >> >> "
                f"/MediaBox [0 0 300 300] /Contents {content_obj_ids[i]} 0 R >>"
            ).encode()
        )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for i in range(page_count):
        stream_content = f"BT /F1 12 Tf 20 250 Td ({text} page {i + 1}) Tj ET".encode("latin-1")
        objects.append(b"<< /Length %d >>\nstream\n" % len(stream_content) + stream_content + b"\nendstream")

    return _assemble_pdf(objects)


def build_blank_pdf(*, page_count: int = 1) -> bytes:
    """A valid PDF with pages but NO text content stream at all — the
    honest stand-in for "scanned PDF with no text layer" (this phase never
    fakes real OCR/rasterization, so a truly empty content stream is the
    correct fixture for the OCR_REQUIRED path).
    """
    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    page_obj_ids = list(range(3, 3 + page_count))
    kids = " ".join(f"{pid} 0 R" for pid in page_obj_ids)
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>".encode())
    for _ in range(page_count):
        objects.append(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] >>")
    return _assemble_pdf(objects)


def build_mixed_pdf(pages: list[str | None]) -> bytes:
    """A valid multi-page PDF where each page independently either has a
    real text content stream (`pages[i]` is a string) or none at all
    (`pages[i]` is `None`) — the fixture for testing that OCR is applied
    per-page: a document with some native-text pages and some scanned pages
    mixed together, exactly like a real multi-page filing where only some
    pages were scanned. See `build_minimal_pdf`'s Latin-1 note — the same
    limitation applies to any non-`None` page here.
    """
    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")

    page_obj_ids = list(range(3, 3 + len(pages)))
    kids = " ".join(f"{pid} 0 R" for pid in page_obj_ids)
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())

    # Only non-blank pages get a content-stream object at all (a blank page
    # has no /Contents key, matching build_blank_pdf) — so content-object
    # IDs must be allocated only for those, in the exact order they'll be
    # emitted below, not one slot per page regardless of whether it's blank.
    font_obj_id = 3 + len(pages)
    content_obj_id_by_index: dict[int, int] = {}
    next_id = font_obj_id + 1
    for i, text in enumerate(pages):
        if text is not None:
            content_obj_id_by_index[i] = next_id
            next_id += 1

    for i, text in enumerate(pages):
        if text is None:
            objects.append(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] >>")
        else:
            objects.append(
                (
                    f"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 {font_obj_id} 0 R >> >> "
                    f"/MediaBox [0 0 300 300] /Contents {content_obj_id_by_index[i]} 0 R >>"
                ).encode()
            )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for text in pages:
        if text is None:
            continue
        stream_content = f"BT /F1 12 Tf 20 250 Td ({text}) Tj ET".encode("latin-1")
        objects.append(b"<< /Length %d >>\nstream\n" % len(stream_content) + stream_content + b"\nendstream")

    return _assemble_pdf(objects)


def _assemble_pdf(objects: list[bytes]) -> bytes:
    out = bytearray()
    out += b"%PDF-1.4\n"
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_offset = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets[1:]:
        out += f"{off:010d} 00000 n \n".encode()
    out += b"trailer\n"
    out += f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode()
    out += b"startxref\n"
    out += f"{xref_offset}\n".encode()
    out += b"%%EOF"
    return bytes(out)


def build_docx(paragraphs: list[str], *, headings: dict[int, str] | None = None) -> bytes:
    """`headings` maps a paragraph index (0-based) to a heading level 1-9;
    that paragraph is added with style "Heading N" instead of plain text.
    """
    import docx

    document = docx.Document()
    headings = headings or {}
    for i, text in enumerate(paragraphs):
        if i in headings:
            document.add_paragraph(text, style=f"Heading {headings[i]}")
        else:
            document.add_paragraph(text)
    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_docx_with_table(rows: list[list[str]]) -> bytes:
    import docx

    document = docx.Document()
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, cell in enumerate(row):
            table.cell(r, c).text = cell
    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_xlsx(sheet_rows: list[list[str]]) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    for row in sheet_rows:
        ws.append(row)
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_zip_bomb_docx() -> bytes:
    """A ZIP with the required Office manifest entry plus one entry with an
    absurd compression ratio — must be rejected by `validate_upload` before
    ever being handed to python-docx.
    """
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", "<Types/>")
        # 50 MB of a single repeated byte compresses to a few KB — a
        # compression ratio far past any real DOCX's, without needing to
        # write anywhere near 50 MB into the ZIP_BOMB uncompressed-size cap.
        zf.writestr("word/document.xml", b"A" * (50 * 1024 * 1024))
    return buf.getvalue()


def build_docx_missing_manifest() -> bytes:
    """A syntactically valid ZIP that is NOT an Office document (no
    `[Content_Types].xml`) — the "corrupted/wrong container" case.
    """
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("hello.txt", "not an office document")
    return buf.getvalue()
